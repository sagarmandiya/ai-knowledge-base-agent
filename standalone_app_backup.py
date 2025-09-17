import streamlit as st
import os
import tempfile
import shutil
from pathlib import Path
from typing import List, Optional
import requests
from io import BytesIO

# LangChain imports
from langchain_community.vectorstores import Chroma
from langchain_community.embeddings import HuggingFaceEmbeddings
from langchain_perplexity import ChatPerplexity
from langchain_core.documents import Document
from langchain.prompts import ChatPromptTemplate
from langchain_core.output_parsers import StrOutputParser
from langchain_core.runnables import RunnablePassthrough
from langchain.text_splitter import RecursiveCharacterTextSplitter

# Document processing imports
import pypdf
from bs4 import BeautifulSoup
import docx

# Page configuration
st.set_page_config(
    page_title="AI Knowledge Base Agent",
    page_icon="🤖",
    layout="wide",
    initial_sidebar_state="expanded"
)

# Custom CSS for better styling
st.markdown("""
<style>
    .main-header {
        font-size: 3rem;
        font-weight: bold;
        text-align: center;
        color: #1f77b4;
        margin-bottom: 2rem;
    }
    .sub-header {
        font-size: 1.5rem;
        font-weight: bold;
        color: #2e8b57;
        margin-top: 2rem;
        margin-bottom: 1rem;
    }
    .info-box {
        background-color: #f0f2f6;
        padding: 1rem;
        border-radius: 0.5rem;
        border-left: 4px solid #1f77b4;
        margin: 1rem 0;
    }
    .success-box {
        background-color: #d4edda;
        padding: 1rem;
        border-radius: 0.5rem;
        border-left: 4px solid #28a745;
        margin: 1rem 0;
    }
    .warning-box {
        background-color: #fff3cd;
        padding: 1rem;
        border-radius: 0.5rem;
        border-left: 4px solid #ffc107;
        margin: 1rem 0;
    }
    .error-box {
        background-color: #f8d7da;
        padding: 1rem;
        border-radius: 0.5rem;
        border-left: 4px solid #dc3545;
        margin: 1rem 0;
    }
</style>
""", unsafe_allow_html=True)

# Initialize session state
if "messages" not in st.session_state:
    st.session_state.messages = []
if "documents" not in st.session_state:
    st.session_state.documents = []
if "vectorstore" not in st.session_state:
    st.session_state.vectorstore = None
if "retriever" not in st.session_state:
    st.session_state.retriever = None
if "rag_chain" not in st.session_state:
    st.session_state.rag_chain = None
if "llm" not in st.session_state:
    st.session_state.llm = None
if "embeddings" not in st.session_state:
    st.session_state.embeddings = None

def initialize_components():
    """Initialize LLM and embeddings"""
    try:
        # Initialize embeddings
        if st.session_state.embeddings is None:
            with st.spinner("Loading embedding model..."):
                st.session_state.embeddings = HuggingFaceEmbeddings(
                    model_name="all-MiniLM-L6-v2"
                )
        
        # Initialize LLM
        if st.session_state.llm is None:
            api_key = st.secrets.get("PERPLEXITY_API_KEY")
            if not api_key:
                st.error("❌ PERPLEXITY_API_KEY not found in secrets. Please add it to your Streamlit secrets.")
                return False
            
            st.session_state.llm = ChatPerplexity(
                model="sonar",
                api_key=api_key,
                temperature=0
            )
        
        return True
    except Exception as e:
        st.error(f"❌ Error initializing components: {str(e)}")
        return False

def create_vectorstore(documents: List[Document]):
    """Create or update vectorstore with documents"""
    try:
        if not documents:
            return None
        
        # Split documents
        text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=500,
            chunk_overlap=50
        )
        splits = text_splitter.split_documents(documents)
        
        # Create vectorstore
        vectorstore = Chroma.from_documents(
            documents=splits,
            embedding=st.session_state.embeddings,
            persist_directory="./chroma_db"
        )
        
        # Create retriever
        retriever = vectorstore.as_retriever(search_kwargs={"k": 3})
        
        # Create RAG chain
        prompt = ChatPromptTemplate.from_template("""
        Answer the question based only on the following context:
        {context}

        Question: {question}
        """)
        
        rag_chain = (
            {"context": retriever, "question": RunnablePassthrough()}
            | prompt
            | st.session_state.llm
            | StrOutputParser()
        )
        
        # Update session state
        st.session_state.vectorstore = vectorstore
        st.session_state.retriever = retriever
        st.session_state.rag_chain = rag_chain
        
        return len(splits)
    except Exception as e:
        st.error(f"❌ Error creating vectorstore: {str(e)}")
        return 0

def load_pdf(file_content: bytes, filename: str) -> List[Document]:
    """Load PDF file and extract text"""
    try:
        pdf_file = BytesIO(file_content)
        pdf_reader = pypdf.PdfReader(pdf_file)
        
        documents = []
        for page_num, page in enumerate(pdf_reader.pages):
            text = page.extract_text()
            if text.strip():
                documents.append(Document(
                    page_content=text,
                    metadata={"source": filename, "page": page_num + 1}
                ))
        
        return documents
    except Exception as e:
        st.error(f"❌ Error loading PDF: {str(e)}")
        return []

def load_text(file_content: bytes, filename: str) -> List[Document]:
    """Load text file"""
    try:
        text = file_content.decode('utf-8')
        return [Document(
            page_content=text,
            metadata={"source": filename}
        )]
    except Exception as e:
        st.error(f"❌ Error loading text file: {str(e)}")
        return []

def load_docx(file_content: bytes, filename: str) -> List[Document]:
    """Load DOCX file"""
    try:
        doc_file = BytesIO(file_content)
        doc = docx.Document(doc_file)
        
        text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
        return [Document(
            page_content=text,
            metadata={"source": filename}
        )]
    except Exception as e:
        st.error(f"❌ Error loading DOCX file: {str(e)}")
        return []

def load_web_content(url: str) -> List[Document]:
    """Load content from web URL"""
    try:
        response = requests.get(url, timeout=10)
        response.raise_for_status()
        
        soup = BeautifulSoup(response.content, 'html.parser')
        
        # Remove script and style elements
        for script in soup(["script", "style"]):
            script.decompose()
        
        # Get text
        text = soup.get_text()
        
        # Clean up text
        lines = (line.strip() for line in text.splitlines())
        chunks = (phrase.strip() for line in lines for phrase in line.split("  "))
        text = ' '.join(chunk for chunk in chunks if chunk)
        
        return [Document(
            page_content=text,
            metadata={"source": url}
        )]
    except Exception as e:
        st.error(f"❌ Error loading web content: {str(e)}")
        return []

def process_uploaded_file(uploaded_file) -> List[Document]:
    """Process uploaded file based on its type"""
    file_content = uploaded_file.getvalue()
    filename = uploaded_file.name
    
    if filename.lower().endswith('.pdf'):
        return load_pdf(file_content, filename)
    elif filename.lower().endswith('.txt'):
        return load_text(file_content, filename)
    elif filename.lower().endswith('.docx'):
        return load_docx(file_content, filename)
    elif filename.lower().endswith('.md'):
        return load_text(file_content, filename)
    else:
        st.error(f"❌ Unsupported file type: {filename}")
        return []

def ask_question(question: str) -> str:
    """Ask question using RAG chain"""
    try:
        if st.session_state.rag_chain is None:
            return "❌ No documents loaded. Please upload documents first."
        
        response = st.session_state.rag_chain.invoke(question)
        return response
    except Exception as e:
        return f"❌ Error processing question: {str(e)}"

def reset_database():
    """Reset the database"""
    st.session_state.documents = []
    st.session_state.vectorstore = None
    st.session_state.retriever = None
    st.session_state.rag_chain = None
    st.session_state.messages = []
    
    # Remove chroma database directory
    import shutil
    if os.path.exists("./chroma_db"):
        shutil.rmtree("./chroma_db")

# Main header
st.markdown('<h1 class="main-header">🤖 AI Knowledge Base Agent</h1>', unsafe_allow_html=True)

# Initialize components
if not initialize_components():
    st.stop()

# Sidebar for document management
with st.sidebar:
    st.markdown("## 📁 Document Management")
    
    # Document count
    doc_count = len(st.session_state.documents)
    st.info(f"📊 Documents loaded: {doc_count}")
    
    st.divider()
    
    # Document Upload
    st.markdown("### 📤 Upload Documents")
    uploaded_file = st.file_uploader(
        "Choose a file",
        type=['pdf', 'txt', 'docx', 'md'],
        help="Upload PDF, TXT, DOCX, or MD files to add to the knowledge base"
    )
    
    if uploaded_file is not None:
        if st.button("📤 Upload Document", type="primary"):
            with st.spinner("Processing document..."):
                documents = process_uploaded_file(uploaded_file)
                if documents:
                    st.session_state.documents.extend(documents)
                    chunks_added = create_vectorstore(st.session_state.documents)
                    if chunks_added > 0:
                        st.success(f"✅ Document uploaded successfully!")
                        st.info(f"📄 {chunks_added} document chunks added")
                        st.rerun()
                    else:
                        st.error("❌ Failed to process document")
    
    st.divider()
    
    # Web Link Processing
    st.markdown("### 🌐 Add Web Content")
    web_url = st.text_input(
        "Enter a web URL",
        placeholder="https://example.com/article",
        help="Add content from a web page to the knowledge base"
    )
    
    if web_url and st.button("🌐 Process Web Link", type="secondary"):
        with st.spinner("Processing web content..."):
            documents = load_web_content(web_url)
            if documents:
                st.session_state.documents.extend(documents)
                chunks_added = create_vectorstore(st.session_state.documents)
                if chunks_added > 0:
                    st.success(f"✅ Web content processed successfully!")
                    st.info(f"📄 {chunks_added} document chunks added")
                    st.rerun()
                else:
                    st.error("❌ Failed to process web content")
    
    st.divider()
    
    # Database Management
    st.markdown("### 🗄️ Database Management")
    if st.button("🗑️ Reset Database", type="secondary"):
        reset_database()
        st.success("✅ Database reset successfully")
        st.rerun()

# Main content area
if doc_count == 0:
    st.markdown('<div class="warning-box">', unsafe_allow_html=True)
    st.markdown("""
    ## ⚠️ No Documents Loaded
    
    Please upload documents or add web content using the sidebar before asking questions.
    
    **Supported file types:**
    - 📄 PDF files
    - 📝 Text files (.txt)
    - 📋 Word documents (.docx)
    - 📖 Markdown files (.md)
    
    **Or add web content:**
    - 🌐 Any publicly accessible web page
    """)
    st.markdown('</div>', unsafe_allow_html=True)
    
    # Example questions (disabled)
    st.markdown("### 💡 Example Questions")
    st.info("Upload documents first to enable these example questions")
    
    example_questions = [
        "What is the main topic of the document?",
        "Can you summarize the key points?",
        "What are the important details mentioned?",
        "How does this relate to the overall context?"
    ]
    
    cols = st.columns(2)
    for i, question in enumerate(example_questions):
        with cols[i % 2]:
            st.button(
                f"❓ {question}",
                disabled=True,
                help="Upload documents first to enable this question"
            )
    
else:
    # Chat interface
    st.markdown("### 💬 Chat with AI Agent")
    
    # Display chat messages
    for message in st.session_state.messages:
        with st.chat_message(message["role"]):
            st.markdown(message["content"])
    
    # Chat input
    if prompt := st.chat_input("Ask a question about your documents..."):
        # Add user message to chat history
        st.session_state.messages.append({"role": "user", "content": prompt})
        with st.chat_message("user"):
            st.markdown(prompt)
        
        # Get AI response
        with st.chat_message("assistant"):
            with st.spinner("Thinking..."):
                response = ask_question(prompt)
                st.markdown(response)
        
        # Add assistant response to chat history
        st.session_state.messages.append({"role": "assistant", "content": response})
    
    # Example questions
    st.markdown("### 💡 Example Questions")
    example_questions = [
        "What is the main topic of the document?",
        "Can you summarize the key points?",
        "What are the important details mentioned?",
        "How does this relate to the overall context?"
    ]
    
    cols = st.columns(2)
    for i, question in enumerate(example_questions):
        with cols[i % 2]:
            if st.button(f"❓ {question}"):
                # Add user message to chat history
                st.session_state.messages.append({"role": "user", "content": question})
                with st.chat_message("user"):
                    st.markdown(question)
                
                # Get AI response
                with st.chat_message("assistant"):
                    with st.spinner("Thinking..."):
                        response = ask_question(question)
                        st.markdown(response)
                
                # Add assistant response to chat history
                st.session_state.messages.append({"role": "assistant", "content": response})
                st.rerun()

# Footer
st.markdown("---")
st.markdown("""
<div style='text-align: center; color: #666;'>
    <p>🤖 AI Knowledge Base Agent | Built with Streamlit & LangChain</p>
    <p>Upload documents or add web content to start chatting with your AI agent!</p>
</div>
""", unsafe_allow_html=True)
